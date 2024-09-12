# A module to convert Persian dates in docx files into Christian dates
# written by Saeed Majidi G., Sep. 11, 2024
# you need the following packages installed
# you need to  type the dates in the following formats ONLY: e.g. 1364/01/20 or 1403/1/20 or 1403/2/2
# you should type the dates with English keyboard.


import os
import re
import jdatetime
from docx import Document
from docx.oxml.ns import qn

# Function to convert Persian date to English date
def convert_persian_to_english_date(persian_date):
    year, month, day = map(int, persian_date.split('/'))
    gregorian_date = jdatetime.date(year, month, day).togregorian()
    return gregorian_date.strftime('%b. %d, %Y')

# Function to replace Persian dates using regex and lambda
def replace_persian_dates(text):
    pattern = r'(\d{4})/(\d{1,2})/(\d{1,2})'
    return re.sub(pattern, lambda m: convert_persian_to_english_date(m.group(0)), text)

# Function to replace Persian dates in a run while keeping formatting
def process_run(run):
    original_text = run.text
    updated_text = replace_persian_dates(original_text)
    
    if original_text != updated_text:
        run.text = updated_text

# Function to process each paragraph
def process_paragraph(paragraph):
    for run in paragraph.runs:
        process_run(run)

# Function to process all elements within a table
def process_table(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                process_paragraph(para)

# Function to process headers and footers
def process_section(section):
    if section.header:
        for para in section.header.paragraphs:
            process_paragraph(para)
    if section.footer:
        for para in section.footer.paragraphs:
            process_paragraph(para)

# Function to process footnotes (if present)
def process_footnotes(doc):
    if hasattr(doc, 'footnotes'):
        footnotes_part = doc.footnotes.part
        if footnotes_part:
            for footnote in footnotes_part.element.body:
                if footnote.tag.endswith('p'):
                    para = footnote
                    process_paragraph(para)

# Function to process comments (if present)
def process_comments(doc):
    if hasattr(doc, 'comments'):
        comments_part = doc.comments.part
        if comments_part:
            for comment in comments_part.element.body:
                if comment.tag.endswith('p'):
                    para = comment
                    process_paragraph(para)

# Function to process every part of the document
def process_docx_file(file_path):
    doc = Document(file_path)

    # Process main body paragraphs
    for para in doc.paragraphs:
        process_paragraph(para)

    # Process tables in the main document
    for table in doc.tables:
        process_table(table)

    # Process headers, footers, and other sections
    for section in doc.sections:
        process_section(section)

    # Process footnotes (if present)
    process_footnotes(doc)

    # Process comments (if present)
    process_comments(doc)

    # Save the updated document
    new_file_path = file_path.replace('.docx', '_updated.docx')
    doc.save(new_file_path)

# Main function to process all .docx files in the current folder
def process_folder():
    folder_path = os.getcwd()
    file_count = 0

    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            file_path = os.path.join(folder_path, filename)
            print(f'Processing {filename}...')
            process_docx_file(file_path)
            file_count += 1

    print(f"Processing complete! {file_count} file(s) processed.")

# Run the script
process_folder()
